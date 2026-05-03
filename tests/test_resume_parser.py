from __future__ import annotations

import docx
from sqlalchemy import select

from db.database import get_session
from db.models import Criterion
from scoring.resume_parser import find_terms, load_taxonomy, parse_resume


def test_word_boundary_excludes_substrings():
    assert find_terms("I am pythonic.", ["python"]) == []
    assert find_terms("I love python and SQL.", ["python"]) == ["python"]


def test_substring_in_other_word_does_not_match():
    assert find_terms("javascript developer", ["java"]) == []
    assert find_terms("javascript developer", ["javascript"]) == ["javascript"]


def test_special_char_terms_match():
    text = "Strong skills in C++, C# and a/b testing experience."
    assert find_terms(text, ["c++"]) == ["c++"]
    assert find_terms(text, ["c#"]) == ["c#"]
    assert find_terms(text, ["a/b testing"]) == ["a/b testing"]


def test_taxonomy_loads_and_has_required_categories():
    tax = load_taxonomy()
    assert "skills" in tax
    assert "roles" in tax
    assert "python" in tax["skills"]["languages"]
    assert "data analyst" in tax["roles"]


def test_parse_resume_extracts_expected_criteria(fresh_db, tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Senior Data Analyst with 5 years of experience.")
    doc.add_paragraph("Tech: Python, SQL, pandas, numpy, scikit-learn.")
    doc.add_paragraph("Built dashboards in Tableau and Streamlit.")
    doc.add_paragraph("AWS, PostgreSQL, Snowflake, Airflow, dbt.")
    doc.add_paragraph("Concepts: ETL, A/B testing, regression, classification.")
    resume_path = tmp_path / "test_resume.docx"
    doc.save(resume_path)

    profile = parse_resume(str(resume_path), "test_profile")
    assert profile.id is not None
    assert profile.resume_filename == "test_resume.docx"
    assert "Python" in profile.resume_raw_text

    with get_session() as session:
        criteria = (
            session.execute(select(Criterion).where(Criterion.profile_id == profile.id))
            .scalars()
            .all()
        )
    terms = {c.term for c in criteria}
    kinds_by_term = {c.term: c.kind for c in criteria}
    weights_by_term = {c.term: c.weight for c in criteria}

    expected_skills = {
        "python",
        "sql",
        "pandas",
        "numpy",
        "scikit-learn",
        "tableau",
        "streamlit",
        "aws",
        "postgresql",
        "snowflake",
        "airflow",
        "dbt",
        "etl",
        "a/b testing",
        "regression",
        "classification",
    }
    assert expected_skills.issubset(terms), f"Missing skills: {expected_skills - terms}"

    assert "data analyst" in terms
    assert kinds_by_term["data analyst"] == "role"
    assert weights_by_term["data analyst"] == 4
    assert kinds_by_term["python"] == "skill"
    assert weights_by_term["python"] == 3


def test_parse_resume_replaces_resume_sourced_criteria(fresh_db, tmp_path):
    """Re-parsing the same profile should not duplicate resume-sourced rows,
    and should preserve manually-added criteria."""
    doc = docx.Document()
    doc.add_paragraph("Python and SQL skills.")
    resume_path = tmp_path / "r.docx"
    doc.save(resume_path)

    profile = parse_resume(str(resume_path), "p1")

    with get_session() as session:
        session.add(
            Criterion(
                profile_id=profile.id,
                term="manual_term",
                kind="keyword",
                weight=5,
                source="manual",
            )
        )
        session.commit()

    parse_resume(str(resume_path), "p1")

    with get_session() as session:
        rows = (
            session.execute(select(Criterion).where(Criterion.profile_id == profile.id))
            .scalars()
            .all()
        )
    terms = [c.term for c in rows]
    assert terms.count("python") == 1
    assert "manual_term" in terms
